# Versión adaptada para Asistente Legal Laboral - Sector Perfumistas
# Carga doc + feedback + logos + descarga Word

from __future__ import annotations

import os
import io
import streamlit as st
from dotenv import load_dotenv
from openai import OpenAI
from pypdf import PdfReader
from docx import Document
from PIL import Image

# Local .env
load_dotenv()

# --- CONFIGURACIÓN DE AVATARES ---
try:
    # He actualizado los nombres de las variables para que coincidan con la temática
    logo_path = "assistant_logo.png" 
    logo_path_user = "user_logo.png"
    USER_AVATAR = Image.open(logo_path_user)
    ASSISTANT_AVATAR = Image.open(logo_path)
except Exception:
    # Emojis actualizados a la temática laboral/perfumería
    USER_AVATAR = "👤"
    ASSISTANT_AVATAR = "🧪" 

API_KEY = os.getenv("OPENAI_API_KEY", "")
VS_ID = os.getenv("OPENAI_VECTOR_STORE_ID", "")
MODEL = os.getenv("OPENAI_MODEL", "gpt-5.1")

# Título de la página actualizado
st.set_page_config(page_title="Asistente Laboral Perfumistas", layout="wide", initial_sidebar_state="expanded")

with st.sidebar:
    st.title("📂 Documentación Laboral")
    archivo = st.file_uploader("Cargar CCT, Recibos o Telegramas para analizar", type=["pdf", "docx"])
    texto_del_archivo = ""
    
    if archivo:
        if archivo.type == "application/pdf":
            reader = PdfReader(archivo)
            texto_del_archivo = "\n".join([p.extract_text() for p in reader.pages if p.extract_text()])
        else:
            doc = Document(archivo)
            texto_del_archivo = "\n".join([p.text for p in doc.paragraphs])
        st.success("Documento cargado con éxito")

col_logo, col_titulo = st.columns([1, 8])
with col_logo:
    # Asegúrate de tener este archivo o cambiarlo por el nombre de tu logo actual
    st.image("logo.png", width=80) 

with col_titulo:
    st.title("Chalk Legal | Laboral")

if not API_KEY or not VS_ID:
    st.error("Faltan credenciales en el archivo .env.")
    st.stop()

client = OpenAI(api_key=API_KEY)

if "messages" not in st.session_state:
    st.session_state.messages = [
        {
            "role": "system",
            "content": (
                "Sos un experto en derecho laboral argentino, especializado en el Convenio Colectivo de Trabajo de Perfumistas. "
                "Tu rol es asesorar sobre escalas salariales, categorías, licencias y normativas específicas del sector. "
                "Respondé en español, citando artículos del CCT o leyes laborales cuando corresponda. "
                "Si no hay soporte documental en el Vector Store sobre una consulta específica, indicalo. "
                "Responder solo en base a los documentos cargados en el VS y legislación laboral vigente."
            ),
        }
    ]

# Render de historial y feedback
for i, m in enumerate(st.session_state.messages):
    if m["role"] in ("user", "assistant"):
        avatar_to_use = USER_AVATAR if m["role"] == "user" else ASSISTANT_AVATAR
        with st.chat_message(m["role"], avatar=avatar_to_use):
            content_to_show = m.get("display_content", m["content"])
            st.markdown(content_to_show)
            
            if m["role"] == "assistant" and i > 0:
                col1, col2 = st.columns([1, 1])
                with col1:
                    key_fb = f"feedback_{i}"
                    fback = st.feedback("thumbs", key=key_fb)
                with col2:
                    doc_download = Document()
                    # Encabezado del documento Word actualizado
                    doc_download.add_heading('Informe Laboral - Sector Perfumistas', 0)
                    doc_download.add_paragraph(m["content"])
                    buffer = io.BytesIO()
                    doc_download.save(buffer)
                    word_data = buffer.getvalue()
                    st.download_button(
                        label="📥 Descargar Informe Laboral",
                        data=word_data,
                        file_name=f"consulta_laboral_perfumistas_{i}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"dl_{i}"
                    )
                if fback == 0:
                    with st.container():
                        comentario = st.text_area("¿En qué podemos precisar mejor la respuesta legal?", key=f"msg_{key_fb}")
                        if st.button("Enviar feedback", key=f"btn_{key_fb}"):
                            st.success("¡Gracias! Analizaremos el caso.")

user_text = st.chat_input("Consulta sobre CCT, escalas o normativa laboral...")

if user_text:
    prompt_final = user_text
    if texto_del_archivo:
        prompt_final = f"DOCUMENTO ADJUNTO DEL EMPLEADO/EMPRESA:\n{texto_del_archivo}\n\nCONSULTA LABORAL: {user_text}"

    st.session_state.messages.append({
        "role": "user", 
        "content": prompt_final, 
        "display_content": user_text
    })
    st.rerun()

# Lógica de respuesta
if len(st.session_state.messages) > 0 and st.session_state.messages[-1]["role"] == "user":
    with st.chat_message("assistant", avatar=ASSISTANT_AVATAR):
        placeholder = st.empty()
        
        api_messages = [
            {"role": m["role"], "content": m["content"]} 
            for m in st.session_state.messages
        ]

        with st.spinner("Consultando CCT de Perfumistas y Ley de Contrato de Trabajo..."):
            resp = client.responses.create(
                model=MODEL,
                input=api_messages,
                tools=[{"type": "file_search", "vector_store_ids": [VS_ID]}],
            )

        answer_text = resp.output_text or "(No se pudo generar una respuesta legal)"
        placeholder.markdown(answer_text)
        st.session_state.messages.append({"role": "assistant", "content": answer_text})

        st.rerun()




